#!/usr/bin/env python3
"""
Convert Markdown to Word Document
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='SimSun', font_size=12, bold=False):
    """Set Chinese font for a run"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold

def convert_md_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""
    doc = Document()
    
    # Set default font for document
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(12)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip HTML comments and image references
        if line.strip().startswith('![') or line.strip().startswith('[') and '](' in line:
            i += 1
            continue
        
        # Skip special annotation lines
        if '【配图' in line or '**【配图' in line:
            i += 1
            continue
        
        # Skip divider lines
        if line.strip() == '---':
            i += 1
            continue
        
        # H1 Title
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=1)
            for run in p.runs:
                set_chinese_font(run, 'Microsoft YaHei', 18, bold=True)
            i += 1
            continue
        
        # H2 Heading
        if line.startswith('## '):
            title = line[3:].strip()
            p = doc.add_heading(title, level=2)
            for run in p.runs:
                set_chinese_font(run, 'Microsoft YaHei', 16, bold=True)
            i += 1
            continue
        
        # H3 Heading
        if line.startswith('### '):
            title = line[4:].strip()
            p = doc.add_heading(title, level=3)
            for run in p.runs:
                set_chinese_font(run, 'Microsoft YaHei', 14, bold=True)
            i += 1
            continue
        
        # Code block
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            set_chinese_font(run, 'Consolas', 10)
            p.paragraph_format.left_indent = Inches(0.5)
            i += 1
            continue
        
        # Table
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            # Parse table header
            headers = [cell.strip() for cell in line.split('|') if cell.strip()]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            # Add header cells
            hdr_cells = table.rows[0].cells
            for j, header in enumerate(headers):
                hdr_cells[j].text = header
                for paragraph in hdr_cells[j].paragraphs:
                    for run in paragraph.runs:
                        set_chinese_font(run, 'Microsoft YaHei', 11, bold=True)
            
            i += 2  # Skip header separator
            
            # Add table rows
            while i < len(lines) and '|' in lines[i] and not lines[i].strip().startswith('#'):
                cells = [cell.strip() for cell in lines[i].split('|') if cell.strip()]
                if cells:
                    row_cells = table.add_row().cells
                    for j, cell_text in enumerate(cells[:len(headers)]):
                        row_cells[j].text = cell_text
                        for paragraph in row_cells[j].paragraphs:
                            for run in paragraph.runs:
                                set_chinese_font(run, 'SimSun', 11)
                i += 1
            continue
        
        # Regular paragraph with formatting
        if line.strip():
            p = doc.add_paragraph()
            
            # Process inline formatting
            parts = re.split(r'(\*\*\*|\*\*|\*|`|~~)', line)
            is_bold = False
            is_italic = False
            is_code = False
            
            for part in parts:
                if part == '***':
                    is_bold = not is_bold
                    is_italic = not is_italic
                elif part == '**':
                    is_bold = not is_bold
                elif part == '*':
                    is_italic = not is_italic
                elif part == '`':
                    is_code = not is_code
                elif part:
                    run = p.add_run(part)
                    if is_code:
                        set_chinese_font(run, 'Consolas', 10)
                    elif is_bold and is_italic:
                        set_chinese_font(run, 'SimSun', 12, bold=True)
                        run.font.italic = True
                    elif is_bold:
                        set_chinese_font(run, 'SimSun', 12, bold=True)
                    elif is_italic:
                        set_chinese_font(run, 'SimSun', 12)
                        run.font.italic = True
                    else:
                        set_chinese_font(run, 'SimSun', 12)
            
            i += 1
        else:
            # Empty line
            i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"Document saved: {docx_file}")

if __name__ == '__main__':
    import sys
    if len(sys.argv) >= 2:
        md_file = sys.argv[1]
        docx_file = sys.argv[2] if len(sys.argv) > 2 else md_file.replace('.md', '.docx')
        convert_md_to_docx(md_file, docx_file)
    else:
        # Default files
        convert_md_to_docx(
            '/Users/cgz/.openclaw/workspace/articles/2026-02-04-AI选股系统完整教程.md',
            '/Users/cgz/.openclaw/workspace/articles/2026-02-04-AI选股系统完整教程.docx'
        )
